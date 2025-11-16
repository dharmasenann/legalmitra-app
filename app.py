import streamlit as st
import os
from pathlib import Path
import google.generativeai as genai
import json
from datetime import datetime
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO

# Set page config - MUST BE FIRST
st.set_page_config(
    page_title="LegalMitra - AI Legal Assistant",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Beautiful Modern CSS with Eye-Friendly Colors
st.markdown("""
<style>
    /* Main App Styling */
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    
    /* Hide Streamlit Branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* Custom Header */
    .custom-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 30px;
        border-radius: 20px;
        margin-bottom: 30px;
        box-shadow: 0 10px 40px rgba(0,0,0,0.1);
        text-align: center;
    }
    
    .custom-header h1 {
        color: white;
        font-size: 48px;
        font-weight: 800;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
    }
    
    .custom-header p {
        color: #e0e7ff;
        font-size: 20px;
        margin: 10px 0 0 0;
    }
    
    /* Card Styling */
    .feature-card {
        background: white;
        padding: 25px;
        border-radius: 15px;
        box-shadow: 0 5px 20px rgba(0,0,0,0.08);
        margin: 15px 0;
        transition: transform 0.3s ease, box-shadow 0.3s ease;
        border-left: 5px solid #667eea;
    }
    
    .feature-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
    }
    
    /* Button Styling */
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 15px 30px;
        border-radius: 12px;
        font-size: 16px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
        width: 100%;
    }
    
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
    }
    
    /* Input Fields */
    .stTextInput>div>div>input,
    .stTextArea>div>div>textarea {
        background: white;
        border: 2px solid #e0e7ff;
        border-radius: 12px;
        padding: 15px;
        font-size: 16px;
        transition: border 0.3s ease;
    }
    
    .stTextInput>div>div>input:focus,
    .stTextArea>div>div>textarea:focus {
        border: 2px solid #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }
    
    /* Sidebar Styling */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #667eea 0%, #764ba2 100%);
    }
    
    [data-testid="stSidebar"] .stMarkdown {
        color: white;
    }
    
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {
        color: white !important;
    }
    
    /* Success/Info/Warning Boxes */
    .stSuccess {
        background: linear-gradient(135deg, #84fab0 0%, #8fd3f4 100%);
        border-radius: 12px;
        padding: 15px;
        border-left: 5px solid #10b981;
    }
    
    .stInfo {
        background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);
        border-radius: 12px;
        padding: 15px;
        border-left: 5px solid #3b82f6;
    }
    
    .stWarning {
        background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
        border-radius: 12px;
        padding: 15px;
        border-left: 5px solid #f59e0b;
    }
    
    /* Metric Cards */
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 25px;
        border-radius: 15px;
        color: white;
        text-align: center;
        box-shadow: 0 5px 20px rgba(102, 126, 234, 0.3);
        margin: 10px 0;
    }
    
    .metric-value {
        font-size: 48px;
        font-weight: 800;
        margin: 10px 0;
    }
    
    .metric-label {
        font-size: 16px;
        opacity: 0.9;
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    
    /* Case Card */
    .case-card {
        background: white;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0 3px 15px rgba(0,0,0,0.08);
        margin: 15px 0;
        border-left: 5px solid #10b981;
        transition: all 0.3s ease;
    }
    
    .case-card:hover {
        box-shadow: 0 5px 25px rgba(0,0,0,0.12);
        transform: translateX(5px);
    }
    
    /* Analysis Result Box */
    .analysis-box {
        background: white;
        padding: 30px;
        border-radius: 15px;
        box-shadow: 0 5px 20px rgba(0,0,0,0.08);
        margin: 20px 0;
    }
    
    .analysis-box h2 {
        color: #667eea;
        border-bottom: 3px solid #667eea;
        padding-bottom: 10px;
        margin-bottom: 20px;
    }
    
    .analysis-box h3 {
        color: #764ba2;
        margin-top: 25px;
    }
    
    /* Expander Styling */
    .streamlit-expanderHeader {
        background: linear-gradient(135deg, #f5f7fa 0%, #e0e7ff 100%);
        border-radius: 10px;
        font-weight: 600;
    }
    
    /* Tab Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: white;
        border-radius: 10px 10px 0 0;
        padding: 15px 25px;
        font-weight: 600;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
    }
    
    /* File Uploader */
    [data-testid="stFileUploader"] {
        background: white;
        border-radius: 15px;
        padding: 20px;
        border: 2px dashed #667eea;
    }
    
    /* Progress Bar */
    .stProgress > div > div > div {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    /* Selectbox */
    .stSelectbox>div>div {
        background: white;
        border-radius: 12px;
    }
    
    /* Responsive Design */
    @media (max-width: 768px) {
        .custom-header h1 {
            font-size: 32px;
        }
        
        .custom-header p {
            font-size: 16px;
        }
        
        .metric-value {
            font-size: 36px;
        }
        
        .stButton>button {
            padding: 12px 20px;
            font-size: 14px;
        }
    }
    
    /* Badge Styling */
    .badge {
        display: inline-block;
        padding: 8px 16px;
        border-radius: 20px;
        font-size: 14px;
        font-weight: 600;
        margin: 5px;
    }
    
    .badge-success {
        background: linear-gradient(135deg, #84fab0 0%, #8fd3f4 100%);
        color: #047857;
    }
    
    .badge-warning {
        background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
        color: #92400e;
    }
    
    .badge-info {
        background: linear-gradient(135deg, #a8edea 0%, #fed6e3 100%);
        color: #1e40af;
    }
    
    /* Loading Animation */
    .stSpinner > div {
        border-top-color: #667eea !important;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'case_history' not in st.session_state:
    st.session_state.case_history = []
if 'legal_context' not in st.session_state:
    st.session_state.legal_context = ""
if 'documents_processed' not in st.session_state:
    st.session_state.documents_processed = False
if 'current_language' not in st.session_state:
    st.session_state.current_language = 'English'
if 'case_id_counter' not in st.session_state:
    st.session_state.case_id_counter = 1
if 'saved_cases_db' not in st.session_state:
    st.session_state.saved_cases_db = {}
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""
if 'api_key_saved' not in st.session_state:
    st.session_state.api_key_saved = False

def process_pdf(pdf_path):
    """Extract text from PDF"""
    try:
        import PyPDF2
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(min(30, len(pdf_reader.pages))):
                text += pdf_reader.pages[page_num].extract_text() + "\n\n"
        return text
    except Exception as e:
        return f"Error: {str(e)}"

def analyze_case(api_key, case_scenario, legal_context, language='English'):
    """Analyze case using Gemini"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        
        lang_instruction = ""
        if language == 'Hindi':
            lang_instruction = "Provide the analysis in Hindi (Devanagari script)."
        elif language == 'Telugu':
            lang_instruction = "Provide the analysis in Telugu (Telugu script)."
        elif language == 'Tamil':
            lang_instruction = "Provide the analysis in Tamil (Tamil script)."
        
        prompt = f"""You are an expert Indian legal analyst. {lang_instruction}

CASE: {case_scenario}

LEGAL REFERENCE: {legal_context if legal_context else "Use IPC, CrPC, Evidence Act knowledge."}

Provide comprehensive analysis:

1. CASE CLASSIFICATION
2. RELEVANT LEGAL PROVISIONS (section numbers)
3. PUNISHMENT DETAILS:
   - Each applicable section
   - Minimum punishment
   - Maximum punishment
   - Cognizable/Non-cognizable
   - Bailable/Non-bailable
4. PROSECUTION ARGUMENTS (4-5 points)
5. DEFENSE ARGUMENTS (4-5 points)
6. KEY LEGAL FACTORS
7. EVIDENCE REQUIREMENTS
8. SIMILAR PRECEDENTS
9. RISK ASSESSMENT

Format clearly with headings."""

        response = model.generate_content(prompt)
        
        case_id = f"CASE-{st.session_state.case_id_counter:06d}"
        st.session_state.case_id_counter += 1
        
        case_data = {
            'case_id': case_id,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'case_scenario': case_scenario,
            'full_scenario': case_scenario,
            'analysis': response.text,
            'language': language,
            'evidence_updates': [],
            'version': 1,  # Track versions for permalink updates
            'last_updated': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        st.session_state.case_history.append(case_data)
        st.session_state.saved_cases_db[case_id] = case_data
        
        return response.text, case_id
    except Exception as e:
        return f"Error: {str(e)}", None

def update_with_evidence(api_key, case_id, original_analysis, new_evidence, language='English'):
    """Update analysis with new evidence"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        
        lang_instruction = ""
        if language == 'Hindi':
            lang_instruction = "Provide update in Hindi."
        elif language == 'Telugu':
            lang_instruction = "Provide update in Telugu."
        elif language == 'Tamil':
            lang_instruction = "Provide update in Tamil."
        
        prompt = f"""{lang_instruction}

ORIGINAL ANALYSIS: {original_analysis}

NEW EVIDENCE: {new_evidence}

Analyze impact:
1. IMPACT ASSESSMENT
2. HOW IT AFFECTS PUNISHMENT
3. EFFECT ON PROSECUTION
4. EFFECT ON DEFENSE
5. UPDATED RISK ASSESSMENT"""

        response = model.generate_content(prompt)
        
        if case_id in st.session_state.saved_cases_db:
            # Update version number
            st.session_state.saved_cases_db[case_id]['version'] = st.session_state.saved_cases_db[case_id].get('version', 1) + 1
            st.session_state.saved_cases_db[case_id]['last_updated'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            st.session_state.saved_cases_db[case_id]['evidence_updates'].append({
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'evidence': new_evidence,
                'impact_analysis': response.text,
                'version': st.session_state.saved_cases_db[case_id]['version']
            })
        
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

def search_precedents(api_key, case_scenario):
    """Search precedents"""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('models/gemini-2.5-flash')
        
        prompt = f"""Find similar Indian legal precedents for: {case_scenario}

List 5-7 relevant precedents with:
1. Case name
2. Year
3. Court
4. Key principle
5. Relevance"""

        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

def load_case_by_id(case_id):
    """Load case by ID"""
    if case_id in st.session_state.saved_cases_db:
        return st.session_state.saved_cases_db[case_id]
    return None

def export_to_pdf(case_scenario, analysis, case_id):
    """Generate PDF report with Unicode support"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4, 
            topMargin=0.5*inch, 
            bottomMargin=0.5*inch,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch
        )
        
        # Register Unicode fonts for Telugu/Hindi support
        try:
            # Try to use system fonts that support Unicode
            pdfmetrics.registerFont(TTFont('Telugu', 'NotoSansTelugu-Regular.ttf'))
            pdfmetrics.registerFont(TTFont('Hindi', 'NotoSansDevanagari-Regular.ttf'))
            font_name = 'Telugu'
        except:
            # Fallback to default font
            font_name = 'Helvetica'
        
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles with Unicode support
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            spaceAfter=20,
            alignment=1,  # Center
            fontName=font_name
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#764ba2'),
            spaceAfter=10,
            spaceBefore=15,
            fontName=font_name,
            leading=20
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            fontName=font_name,
            leading=16,
            spaceAfter=8
        )
        
        # Add gradient header box
        header_data = [[Paragraph("Legal Case Analysis Report", title_style)]]
        header_table = Table(header_data, colWidths=[6.5*inch])
        header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#667eea')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 20),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 20))
        
        # Case Information Box
        info_data = [
            [Paragraph(f"<b>Case ID:</b> {case_id}", normal_style),
             Paragraph(f"<b>Date:</b> {datetime.now().strftime('%d-%m-%Y %H:%M')}", normal_style)]
        ]
        info_table = Table(info_data, colWidths=[3.25*inch, 3.25*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f0f9ff')),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#667eea')),
            ('TOPPADDING', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))
        
        # Case Scenario Section
        story.append(Paragraph("Case Scenario", heading_style))
        # Clean and encode text properly
        clean_scenario = case_scenario.encode('utf-8', errors='ignore').decode('utf-8')
        for para in clean_scenario.split('\n'):
            if para.strip():
                story.append(Paragraph(para, normal_style))
        story.append(Spacer(1, 20))
        
        # Analysis Section
        story.append(Paragraph("Legal Analysis", heading_style))
        
        # Clean and format analysis text
        clean_analysis = analysis.encode('utf-8', errors='ignore').decode('utf-8')
        
        # Split by lines and format
        for line in clean_analysis.split('\n'):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            # Check if it's a heading (starts with ** or #)
            if line.startswith('**') and line.endswith('**'):
                heading_text = line.replace('**', '').strip()
                story.append(Paragraph(heading_text, heading_style))
            elif line.startswith('#'):
                heading_text = line.replace('#', '').strip()
                story.append(Paragraph(heading_text, heading_style))
            elif line.startswith('*') or line.startswith('-'):
                # Bullet point
                bullet_text = line.lstrip('*- ').strip()
                bullet_para = Paragraph(f"• {bullet_text}", normal_style)
                story.append(bullet_para)
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                # Numbered list
                story.append(Paragraph(line, normal_style))
            else:
                # Regular paragraph
                story.append(Paragraph(line, normal_style))
        
        story.append(Spacer(1, 30))
        
        # Footer
        footer_data = [[Paragraph("<i>Generated by LegalMitra - Your AI Legal Assistant | Powered by Google Gemini 2.5 Flash</i>", 
                                  ParagraphStyle('Footer', parent=styles['Normal'], 
                                               fontSize=9, textColor=colors.grey, alignment=1))]]
        footer_table = Table(footer_data, colWidths=[6.5*inch])
        footer_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ]))
        story.append(footer_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"PDF Error: {str(e)}")
        st.info("💡 For Telugu/Hindi support, the system needs Unicode fonts. PDF generated with available fonts.")
        return None

def export_to_word(case_scenario, analysis, case_id):
    """Generate Word document with full Unicode support"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # Set default font to support Unicode (especially Telugu/Hindi/Tamil)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Noto Sans'  # Universal font for multiple scripts
        font.size = Pt(12)
        
        # Fallback fonts for better compatibility
        rPr = style.element.rPr
        rFonts = rPr.rFonts
        rFonts.set(qn('w:ascii'), 'Noto Sans')
        rFonts.set(qn('w:hAnsi'), 'Noto Sans')
        rFonts.set(qn('w:eastAsia'), 'Noto Sans')
        rFonts.set(qn('w:cs'), 'Noto Sans')  # Complex scripts
        
        # Add gradient-like header
        header = doc.add_heading('', 0)
        header_run = header.add_run('⚖️ Legal Case Analysis Report')
        header_run.font.size = Pt(24)
        header_run.font.color.rgb = RGBColor(102, 126, 234)
        header_run.bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some space
        doc.add_paragraph()
        
        # Case Information
        info_para = doc.add_paragraph()
        info_para.add_run(f"Case ID: {case_id}\n").bold = True
        info_para.add_run(f"Date: {datetime.now().strftime('%d-%m-%Y %H:%M')}\n").bold = True
        info_para.add_run(f"Report Generated by LegalMitra").bold = True
        
        doc.add_paragraph()
        
        # Case Scenario
        doc.add_heading('Case Scenario', level=1)
        scenario_heading = doc.paragraphs[-1]
        scenario_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)
        
        # Add scenario with proper encoding
        scenario_para = doc.add_paragraph(case_scenario)
        for run in scenario_para.runs:
            run.font.name = 'Noto Sans'
        
        doc.add_paragraph()
        
        # Analysis
        doc.add_heading('Legal Analysis', level=1)
        analysis_heading = doc.paragraphs[-1]
        analysis_heading.runs[0].font.color.rgb = RGBColor(118, 75, 162)
        
        # Parse and format analysis
        for line in analysis.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Headings
            if line.startswith('**') and line.endswith('**'):
                heading_text = line.replace('**', '').strip()
                h = doc.add_heading(heading_text, level=2)
                h.runs[0].font.color.rgb = RGBColor(102, 126, 234)
                for run in h.runs:
                    run.font.name = 'Noto Sans'
            elif line.startswith('#'):
                heading_text = line.replace('#', '').strip()
                h = doc.add_heading(heading_text, level=2)
                h.runs[0].font.color.rgb = RGBColor(102, 126, 234)
                for run in h.runs:
                    run.font.name = 'Noto Sans'
            # Bullet points
            elif line.startswith('*') or line.startswith('-'):
                bullet_text = line.lstrip('*- ').strip()
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                for run in para.runs:
                    run.font.name = 'Noto Sans'
            # Numbered lists
            elif line[0].isdigit() and '.' in line[:3]:
                para = doc.add_paragraph(line, style='List Number')
                for run in para.runs:
                    run.font.name = 'Noto Sans'
            # Regular paragraphs
            else:
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.name = 'Noto Sans'
        
        # Footer
        doc.add_paragraph()
        doc.add_paragraph()
        footer = doc.add_paragraph("Generated by LegalMitra | Powered by Google Gemini 2.5 Flash")
        footer.runs[0].italic = True
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Word Error: {str(e)}")
        return None

def generate_visual_charts(analysis_text):
    """Generate visualization data from analysis"""
    # Parse analysis to extract data for charts
    risk_data = {
        'Prosecution Strength': 70,
        'Defense Strength': 30
    }
    
    evidence_data = {
        'Documentary': 40,
        'Witness Testimony': 30,
        'Physical Evidence': 20,
        'Digital Evidence': 10
    }
    
    timeline_data = pd.DataFrame({
        'Stage': ['Incident', 'FIR Filed', 'Investigation', 'Arrest', 'Charge Sheet', 'Trial'],
        'Status': ['Complete', 'Complete', 'Complete', 'Complete', 'Pending', 'Upcoming'],
        'Progress': [100, 100, 100, 100, 50, 0]
    })
    
    return risk_data, evidence_data, timeline_data

def get_case_permalink(case_id, version=1):
    """Generate unique permalink with version"""
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    return f"legalmitra://case/{case_id}/v{version}/{timestamp}"

def load_case_by_id(case_id):
    """Load case from database by ID"""
    if case_id in st.session_state.saved_cases_db:
        return st.session_state.saved_cases_db[case_id]
    return None

def main():
    # Custom Header
    st.markdown("""
    <div class="custom-header">
        <h1>⚖️ LegalMitra</h1>
        <p>Your AI-Powered Indian Legal Assistant</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar
    with st.sidebar:
        st.markdown("## ⚙️ Settings")
        
        # API Key Management
        if not st.session_state.api_key_saved:
            st.markdown("### 🔑 Setup API Key")
            api_key_input = st.text_input("Enter Gemini API Key", type="password",
                                          help="Enter once - will be saved for this session")
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Save Key", use_container_width=True):
                    if api_key_input:
                        st.session_state.api_key = api_key_input
                        st.session_state.api_key_saved = True
                        st.success("✅ API Key Saved!")
                        st.rerun()
                    else:
                        st.error("Please enter API key")
            
            with col2:
                if st.button("❓ Get Key", use_container_width=True):
                    st.info("Visit: aistudio.google.com/app/apikey")
        else:
            st.success("✅ API Key Active")
            st.caption(f"Key: {st.session_state.api_key[:15]}...")
            
            if st.button("🔄 Change API Key", use_container_width=True):
                st.session_state.api_key_saved = False
                st.session_state.api_key = ""
                st.rerun()
        
        api_key = st.session_state.api_key
        
        st.markdown("---")
        
        # Language
        st.markdown("## 🌐 Language")
        language = st.selectbox("Select Language", ['English', 'Hindi', 'Telugu', 'Tamil'], label_visibility="collapsed")
        st.session_state.current_language = language
        
        st.markdown("---")
        
        # Features
        st.markdown("## ✨ Features")
        feature = st.radio("Select Feature", [
            "📝 Analyze Case",
            "📚 Case History",
            "🔍 Find Precedents",
            "📊 Visual Reports",
            "💾 Export"
        ], label_visibility="collapsed")
        
        st.markdown("---")
        
        # Documents
        st.markdown("## 📄 Documents")
        uploaded_files = st.file_uploader("Upload Legal PDFs", 
                                         type=['pdf', 'txt'],
                                         accept_multiple_files=True,
                                         label_visibility="collapsed")
        
        if uploaded_files and api_key:
            documents_path = Path("./legal_documents")
            documents_path.mkdir(exist_ok=True)
            
            for uploaded_file in uploaded_files:
                file_path = documents_path / uploaded_file.name
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
            
            if st.button("📚 Process", use_container_width=True):
                with st.spinner("Processing..."):
                    all_text = ""
                    for pdf_file in documents_path.glob("*.pdf"):
                        text = process_pdf(pdf_file)
                        all_text += text
                    st.session_state.legal_context = all_text[:20000]
                    st.session_state.documents_processed = True
                    st.success("✅ Done!")
        
        st.markdown("---")
        
        # Stats
        st.markdown("## 📊 Statistics")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-value">{len(st.session_state.case_history)}</div>
                <div class="metric-label">Cases</div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            docs_count = "✓" if st.session_state.documents_processed else "0"
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-value">{docs_count}</div>
                <div class="metric-label">Docs</div>
            </div>
            """, unsafe_allow_html=True)
    
    # Main Content
    if not api_key:
        st.markdown("""
        <div class="feature-card">
            <h2>👋 Welcome to LegalMitra</h2>
            <p>Your trusted AI-powered Indian legal research assistant.</p>
            <br>
            <h3>🚀 Features:</h3>
            <ul>
                <li>✅ AI-Powered Case Analysis</li>
                <li>✅ Multi-language Support (English/Hindi/Telugu)</li>
                <li>✅ Precedent Search</li>
                <li>✅ Punishment Details & Ranges</li>
                <li>✅ Evidence Impact Analysis</li>
                <li>✅ Case History & Permalinks</li>
                <li>✅ Export Reports (PDF/Word)</li>
                <li>✅ 100% FREE with Gemini AI</li>
            </ul>
            <br>
            <p><strong>👈 Enter your API key in the sidebar to begin!</strong></p>
            <p><em>Mitra means "Friend" in Sanskrit - Your trusted legal companion</em></p>
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Feature: Analyze Case
    if feature == "📝 Analyze Case":
        # Load case by ID
        col1, col2 = st.columns([3, 1])
        with col1:
            case_id_input = st.text_input("🔗 Load Case by ID", 
                                          placeholder="e.g., CASE-000001")
        with col2:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("📂 Load", use_container_width=True):
                if case_id_input:
                    loaded_case = load_case_by_id(case_id_input)
                    if loaded_case:
                        st.markdown(f"""
                        <div class="feature-card">
                            <h3>✅ Case Loaded: {case_id_input}</h3>
                            <p><strong>Date:</strong> {loaded_case['timestamp']}</p>
                            <p><strong>Language:</strong> {loaded_case['language']}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        st.markdown("""<div class="analysis-box">""", unsafe_allow_html=True)
                        st.markdown(f"**Scenario:** {loaded_case['full_scenario']}")
                        st.markdown("---")
                        st.markdown(loaded_case['analysis'])
                        st.markdown("""</div>""", unsafe_allow_html=True)
                        
                        if loaded_case['evidence_updates']:
                            st.markdown("### 🔎 Evidence History")
                            for idx, update in enumerate(loaded_case['evidence_updates'], 1):
                                with st.expander(f"Update #{idx} - {update['timestamp']}"):
                                    st.markdown(f"**Evidence:** {update['evidence']}")
                                    st.markdown(f"**Impact:** {update['impact_analysis']}")
                    else:
                        st.error("❌ Case not found!")
        
        st.markdown("---")
        
        # New Case Analysis
        st.markdown("### 📝 New Case Analysis")
        case_scenario = st.text_area("Enter case details:", 
                                     height=200,
                                     placeholder="Describe the case with all relevant facts, dates, parties, and circumstances...")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("🔍 Analyze", type="primary", use_container_width=True):
                if case_scenario:
                    with st.spinner("🤖 AI is analyzing..."):
                        analysis, case_id = analyze_case(api_key, case_scenario,
                                                        st.session_state.legal_context,
                                                        language)
                        st.session_state.analysis_result = analysis
                        st.session_state.current_case_id = case_id
                        st.rerun()
        
        with col2:
            if st.button("📖 Precedents", use_container_width=True):
                if case_scenario:
                    with st.spinner("Searching..."):
                        precedents = search_precedents(api_key, case_scenario)
                        st.markdown("""<div class="analysis-box">""", unsafe_allow_html=True)
                        st.markdown("### 📚 Similar Precedents")
                        st.markdown(precedents)
                        st.markdown("""</div>""", unsafe_allow_html=True)
        
        with col3:
            if st.session_state.analysis_result:
                if st.button("📊 Visualize", use_container_width=True):
                    st.session_state['navigate_to_visual'] = True
                    st.info("👉 Please go to '📊 Visual Reports' section in the sidebar to view charts!")
        
        # Display Analysis
        if st.session_state.analysis_result:
            st.markdown("---")
            
            if hasattr(st.session_state, 'current_case_id'):
                col1, col2 = st.columns(2)
                with col1:
                    st.success(f"🔗 **Case ID:** {st.session_state.current_case_id}")
                with col2:
                    case_version = st.session_state.saved_cases_db.get(st.session_state.current_case_id, {}).get('version', 1)
                    permalink = get_case_permalink(st.session_state.current_case_id, case_version)
                    st.info(f"**Permalink:** `{permalink}`")
            
            st.markdown("""<div class="analysis-box">""", unsafe_allow_html=True)
            st.markdown(st.session_state.analysis_result)
            st.markdown("""</div>""", unsafe_allow_html=True)
            
            # Evidence Addition
            st.markdown("---")
            st.markdown("### 🔎 Add Evidence/Document")
            
            new_evidence = st.text_area("Enter new evidence or witness statement:",
                                       height=150,
                                       placeholder="Example: New CCTV footage shows...")
            
            if st.button("🔄 Update Analysis", type="secondary"):
                if new_evidence and hasattr(st.session_state, 'current_case_id'):
                    with st.spinner("🤖 Analyzing evidence..."):
                        update = update_with_evidence(api_key,
                                                     st.session_state.current_case_id,
                                                     st.session_state.analysis_result,
                                                     new_evidence,
                                                     language)
                        
                        st.success("✅ Evidence analyzed!")
                        st.markdown("""<div class="analysis-box">""", unsafe_allow_html=True)
                        st.markdown("### 🔄 Impact Analysis")
                        st.markdown(update)
                        st.markdown("""</div>""", unsafe_allow_html=True)
                        
                        st.info("💡 Add more evidence by entering new details above!")
    
    # Feature: Case History
    elif feature == "📚 Case History":
        st.markdown("## 📚 Case History")
        
        if st.session_state.case_history:
            st.markdown(f"**Total Cases:** {len(st.session_state.case_history)}")
            
            for case in reversed(st.session_state.case_history):
                case_version = case.get('version', 1)
                permalink = get_case_permalink(case['case_id'], case_version)
                last_updated = case.get('last_updated', case['timestamp'])
                
                st.markdown(f"""
                <div class="case-card">
                    <h3>{case['case_id']} <span style="color: #667eea;">v{case_version}</span></h3>
                    <p><strong>📅 Created:</strong> {case['timestamp']}</p>
                    <p><strong>🔄 Last Updated:</strong> {last_updated}</p>
                    <p><strong>🌐 Language:</strong> {case['language']}</p>
                    <p><strong>🔗 Permalink:</strong> <code>{permalink}</code></p>
                    <p><strong>📋 Scenario:</strong> {case['case_scenario'][:150]}...</p>
                    {f'<p><strong>🔎 Evidence Updates:</strong> {len(case["evidence_updates"])}</p>' if case['evidence_updates'] else ''}
                </div>
                """, unsafe_allow_html=True)
                
                with st.expander("📖 View Details"):
                    st.markdown("### Analysis:")
                    st.markdown(case['analysis'])
                    
                    if case['evidence_updates']:
                        st.markdown("### Evidence Updates:")
                        for idx, update in enumerate(case['evidence_updates'], 1):
                            st.markdown(f"**Update #{idx}** - {update['timestamp']}")
                            st.markdown(f"*Evidence:* {update['evidence']}")
                            st.markdown(f"*Impact:* {update['impact_analysis']}")
                            st.markdown("---")
        else:
            st.info("📝 No cases yet. Start by analyzing a case!")
    
    # Feature: Precedents
    elif feature == "🔍 Find Precedents":
        st.markdown("## 🔍 Legal Precedent Search")
        
        search_query = st.text_area("Describe the legal issue:",
                                    height=150)
        
        if st.button("🔎 Search Database", type="primary"):
            if search_query:
                with st.spinner("🔍 Searching Indian case law..."):
                    results = search_precedents(api_key, search_query)
                    st.markdown("""<div class="analysis-box">""", unsafe_allow_html=True)
                    st.markdown("### 📚 Relevant Precedents")
                    st.markdown(results)
                    st.markdown("""</div>""", unsafe_allow_html=True)
    
    # Feature: Visual Reports
    elif feature == "📊 Visual Reports":
        st.markdown("## 📊 Visual Analytics")
        
        if st.session_state.analysis_result and hasattr(st.session_state, 'current_case_id'):
            # Get case data
            case_id = st.session_state.current_case_id
            case_data = st.session_state.saved_cases_db.get(case_id)
            
            if case_data:
                st.success(f"📊 Analyzing Case: {case_id}")
                
                # User Input Section for Case Progress
                st.markdown("### ⚙️ Configure Case Progress")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**📋 Select Completed Stages:**")
                    incident_done = st.checkbox("Incident Occurred", value=True, key="incident")
                    fir_done = st.checkbox("FIR Filed", value=True, key="fir")
                    investigation_done = st.checkbox("Investigation", value=False, key="investigation")
                    arrest_done = st.checkbox("Arrest Made", value=False, key="arrest")
                    chargesheet_done = st.checkbox("Charge Sheet Filed", value=False, key="chargesheet")
                    trial_done = st.checkbox("Trial Commenced", value=False, key="trial")
                
                with col2:
                    st.markdown("**⚖️ Case Assessment:**")
                    prosecution_strength = st.slider("Prosecution Strength", 0, 100, 70, key="pros_strength", help="Set prosecution strength percentage")
                    defense_strength = st.slider("Defense Strength", 0, 100, 30, key="def_strength", help="Set defense strength percentage")
                    
                    st.markdown("**📂 Evidence Distribution:**")
                    doc_evidence = st.slider("Documentary Evidence", 0, 100, 40, key="doc_ev", help="Documentary evidence percentage")
                    witness_evidence = st.slider("Witness Testimony", 0, 100, 30, key="wit_ev", help="Witness testimony percentage")
                    physical_evidence = st.slider("Physical Evidence", 0, 100, 20, key="phy_ev", help="Physical evidence percentage")
                    digital_evidence = st.slider("Digital Evidence", 0, 100, 10, key="dig_ev", help="Digital evidence percentage")
                
                if st.button("🔄 Generate Visual Reports", type="primary", use_container_width=True):
                    # Create timeline data based on user input
                    timeline_data = pd.DataFrame({
                        'Stage': ['Incident', 'FIR Filed', 'Investigation', 'Arrest', 'Charge Sheet', 'Trial'],
                        'Status': [
                            'Complete' if incident_done else 'Pending',
                            'Complete' if fir_done else 'Pending',
                            'Complete' if investigation_done else 'Pending',
                            'Complete' if arrest_done else 'Pending',
                            'Complete' if chargesheet_done else 'Pending',
                            'Complete' if trial_done else 'Pending'
                        ],
                        'Progress': [
                            100 if incident_done else 0,
                            100 if fir_done else 0,
                            100 if investigation_done else 0,
                            100 if arrest_done else 0,
                            100 if chargesheet_done else 0,
                            100 if trial_done else 0
                        ]
                    })
                    
                    risk_data = {
                        'Prosecution Strength': prosecution_strength,
                        'Defense Strength': defense_strength
                    }
                    
                    evidence_data = {
                        'Documentary': doc_evidence,
                        'Witness Testimony': witness_evidence,
                        'Physical Evidence': physical_evidence,
                        'Digital Evidence': digital_evidence
                    }
                    
                    st.markdown("---")
                    
                    # Case Strength Assessment
                    st.markdown("### 🎯 Case Strength Assessment")
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        fig1 = go.Figure(data=[
                            go.Bar(
                                x=list(risk_data.keys()),
                                y=list(risk_data.values()),
                                marker_color=['#ef4444', '#22c55e'],
                                text=[f"{v}%" for v in risk_data.values()],
                                textposition='auto',
                            )
                        ])
                        fig1.update_layout(
                            title="Prosecution vs Defense Strength",
                            yaxis_title="Strength (%)",
                            plot_bgcolor='rgba(0,0,0,0)',
                            paper_bgcolor='rgba(0,0,0,0)',
                            font=dict(color='#1f2937'),
                            height=400
                        )
                        st.plotly_chart(fig1, use_container_width=True)
                    
                    with col2:
                        fig2 = px.pie(
                            values=list(evidence_data.values()),
                            names=list(evidence_data.keys()),
                            title="Evidence Distribution",
                            color_discrete_sequence=px.colors.sequential.Purples,
                            hole=0.3
                        )
                        fig2.update_traces(textposition='inside', textinfo='percent+label')
                        fig2.update_layout(
                            plot_bgcolor='rgba(0,0,0,0)',
                            paper_bgcolor='rgba(0,0,0,0)',
                            font=dict(color='#1f2937'),
                            height=400
                        )
                        st.plotly_chart(fig2, use_container_width=True)
                    
                    # Case Timeline
                    st.markdown("### 📅 Case Progress Timeline")
                    fig3 = go.Figure(data=[
                        go.Bar(
                            x=timeline_data['Stage'],
                            y=timeline_data['Progress'],
                            marker_color=['#10b981' if x == 100 else '#9ca3af' for x in timeline_data['Progress']],
                            text=[f"{x}%" for x in timeline_data['Progress']],
                            textposition='auto',
                        )
                    ])
                    fig3.update_layout(
                        title="Case Progress by Stage",
                        yaxis_title="Completion (%)",
                        plot_bgcolor='rgba(0,0,0,0)',
                        paper_bgcolor='rgba(0,0,0,0)',
                        font=dict(color='#1f2937'),
                        height=400
                    )
                    st.plotly_chart(fig3, use_container_width=True)
                    
                    # Risk Assessment Gauge
                    st.markdown("### ⚠️ Overall Risk Assessment")
                    col1, col2, col3 = st.columns(3)
                    
                    # Calculate risk levels
                    pros_risk = "High" if prosecution_strength > 60 else "Medium" if prosecution_strength > 30 else "Low"
                    pros_color = "#ef4444" if pros_risk == "High" else "#f59e0b" if pros_risk == "Medium" else "#10b981"
                    
                    def_strength_level = "High" if defense_strength > 60 else "Medium" if defense_strength > 30 else "Low"
                    def_color = "#10b981" if def_strength_level == "High" else "#f59e0b" if def_strength_level == "Medium" else "#ef4444"
                    
                    completed_stages = sum([incident_done, fir_done, investigation_done, arrest_done, chargesheet_done, trial_done])
                    complexity = "High" if completed_stages > 4 else "Medium" if completed_stages > 2 else "Low"
                    complex_color = "#ef4444" if complexity == "High" else "#f59e0b" if complexity == "Medium" else "#10b981"
                    
                    with col1:
                        st.markdown(f"""
                        <div class="metric-card" style="background: linear-gradient(135deg, {pros_color} 0%, {pros_color} 100%);">
                            <div class="metric-label">Prosecution Risk</div>
                            <div class="metric-value">{pros_risk}</div>
                            <div class="metric-label">{prosecution_strength}%</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col2:
                        st.markdown(f"""
                        <div class="metric-card" style="background: linear-gradient(135deg, {complex_color} 0%, {complex_color} 100%);">
                            <div class="metric-label">Case Complexity</div>
                            <div class="metric-value">{complexity}</div>
                            <div class="metric-label">{completed_stages}/6 Stages</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    with col3:
                        st.markdown(f"""
                        <div class="metric-card" style="background: linear-gradient(135deg, {def_color} 0%, {def_color} 100%);">
                            <div class="metric-label">Defense Strength</div>
                            <div class="metric-value">{def_strength_level}</div>
                            <div class="metric-label">{defense_strength}%</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Data Table
                    st.markdown("### 📊 Case Progress Details")
                    st.dataframe(timeline_data, use_container_width=True, hide_index=True)
                    
                    # Summary
                    st.markdown("### 📝 Analysis Summary")
                    st.markdown(f"""
                    <div class="analysis-box">
                        <p><strong>Case ID:</strong> {case_id}</p>
                        <p><strong>Completed Stages:</strong> {completed_stages} out of 6</p>
                        <p><strong>Prosecution Advantage:</strong> {prosecution_strength - defense_strength:+d}%</p>
                        <p><strong>Total Evidence:</strong> {doc_evidence + witness_evidence + physical_evidence + digital_evidence}%</p>
                        <p><strong>Strongest Evidence:</strong> {max(evidence_data, key=evidence_data.get)} ({max(evidence_data.values())}%)</p>
                    </div>
                    """, unsafe_allow_html=True)
            else:
                st.error("❌ Case data not found!")
        else:
            st.info("📝 Analyze a case first to see visual reports!")
            st.markdown("""
            <div class="feature-card">
                <h3>📊 How to Use Visual Reports</h3>
                <ol>
                    <li>Go to "Analyze Case" and analyze a case</li>
                    <li>Come back to "Visual Reports"</li>
                    <li>Configure the case progress stages and percentages</li>
                    <li>Click "Generate Visual Reports"</li>
                    <li>View interactive charts and analysis</li>
                </ol>
            </div>
            """, unsafe_allow_html=True)
    
    # Feature: Export
    elif feature == "💾 Export":
        st.markdown("## 💾 Export Reports")
        
        # Font installation info box
        st.markdown("""
        <div class="feature-card" style="border-left: 5px solid #f59e0b;">
            <h4>📝 Important: Telugu/Hindi/Tamil Font Requirements</h4>
            <p>For perfect Telugu/Hindi/Tamil display in Word documents, install <strong>Noto Sans</strong> font family.</p>
            <p>👉 Download from: <a href="https://fonts.google.com/noto" target="_blank">Google Fonts - Noto</a></p>
            <p>See detailed installation guide below ⬇️</p>
        </div>
        """, unsafe_allow_html=True)
        
        if st.session_state.analysis_result and hasattr(st.session_state, 'current_case_id'):
            st.success("✅ Analysis available for export!")
            
            # Get case data
            case_id = st.session_state.current_case_id
            case_data = st.session_state.saved_cases_db.get(case_id)
            
            if case_data:
                case_scenario = case_data['full_scenario']
                analysis = case_data['analysis']
                case_language = case_data.get('language', 'English')
                
                # Language-specific recommendation
                if case_language in ['Telugu', 'Hindi', 'Tamil']:
                    st.info(f"💡 **Recommendation:** For {case_language} text, Word format preserves all characters perfectly!")
                    
                    # Font installation guide
                    with st.expander("📤 Need Telugu/Hindi/Tamil fonts? Click here for installation guide"):
                        st.markdown("""
                        ### 📥 Install Required Fonts for Perfect Display
                        
                        **For Telugu Text:**
                        - **Noto Sans Telugu** (Recommended)
                          - Download: [Google Fonts - Noto Sans Telugu](https://fonts.google.com/noto/specimen/Noto+Sans+Telugu)
                          - Direct Download: [GitHub - Noto Fonts](https://github.com/notofonts/telugu/releases)
                        
                        **For Hindi Text:**
                        - **Noto Sans Devanagari** (Recommended)
                          - Download: [Google Fonts - Noto Sans Devanagari](https://fonts.google.com/noto/specimen/Noto+Sans+Devanagari)
                        
                        **For Tamil Text:**
                        - **Noto Sans Tamil** (Recommended)
                          - Download: [Google Fonts - Noto Sans Tamil](https://fonts.google.com/noto/specimen/Noto+Sans+Tamil)
                          - Direct Download: [GitHub - Noto Fonts](https://github.com/notofonts/tamil/releases)
                        
                        **All-in-One Solution:**
                        - **Noto Sans** (Universal - Supports all scripts)
                          - Download: [Google Fonts - Noto Sans](https://fonts.google.com/noto)
                        
                        **How to Install:**
                        1. Download the font file (.ttf or .otf)
                        2. **Windows:** Right-click → Install / Double-click the font file
                        3. **Mac:** Double-click → Install Font
                        4. **Linux:** Copy to `~/.fonts/` or use Font Manager
                        5. Restart Microsoft Word
                        6. Open the downloaded .docx file - All languages will display perfectly!
                        
                        **Alternative Fonts (if Noto is not available):**
                        - Telugu: Pothana2000, Vani, Gautami
                        - Hindi: Mangal, Lohit Devanagari, Aparajita
                        - Tamil: Lohit Tamil, TSC Tamil, Bamini
                        
                        ✅ **Note:** The Word document is already configured to use Noto Sans fonts. 
                        You just need to install the font on your computer to view it correctly!
                        """)
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("""
                    <div class="feature-card">
                        <h3>📝 Word Export (Recommended)</h3>
                        <p><strong>✅ Perfect for Telugu & Hindi!</strong></p>
                        <p>Full Unicode support - all languages display correctly</p>
                        <p>Editable format, professional layout</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button("📝 Generate Word Document", use_container_width=True, type="primary"):
                        with st.spinner("📝 Generating Word document..."):
                            word_buffer = export_to_word(case_scenario, analysis, case_id)
                            
                            if word_buffer:
                                st.download_button(
                                    label="⬇️ Download Word Report (.docx)",
                                    data=word_buffer,
                                    file_name=f"LegalMitra_Analysis_{case_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                                st.success("✅ Word document ready! All Telugu/Hindi text will display perfectly!")
                
                with col2:
                    st.markdown("""
                    <div class="feature-card">
                        <h3>📄 PDF Export</h3>
                        <p><strong>Best for English text</strong></p>
                        <p>Professional PDF format</p>
                        <p>Note: Telugu/Hindi may not display correctly</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    if st.button("📄 Generate PDF", use_container_width=True):
                        if case_language in ['Telugu', 'Hindi', 'Tamil']:
                            st.warning(f"⚠️ Note: PDF may not display {case_language} characters correctly. Word format is recommended!")
                        
                        with st.spinner("📄 Generating PDF..."):
                            pdf_buffer = export_to_pdf(case_scenario, analysis, case_id)
                            
                            if pdf_buffer:
                                st.download_button(
                                    label="⬇️ Download PDF Report",
                                    data=pdf_buffer,
                                    file_name=f"LegalMitra_Analysis_{case_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )
                                st.success("✅ PDF ready for download!")
                
                # Export options comparison
                st.markdown("---")
                st.markdown("### 📊 Format Comparison")
                
                comparison_data = {
                    'Feature': ['Telugu Support', 'Hindi Support', 'Tamil Support', 'English Support', 'Editable', 'File Size', 'Best For'],
                    'Word (.docx)': ['✅ Perfect', '✅ Perfect', '✅ Perfect', '✅ Perfect', '✅ Yes', 'Medium', 'All Languages'],
                    'PDF': ['❌ Limited', '❌ Limited', '❌ Limited', '✅ Perfect', '❌ No', 'Small', 'English Only']
                }
                
                import pandas as pd
                df = pd.DataFrame(comparison_data)
                st.dataframe(df, use_container_width=True, hide_index=True)
                
                # Preview
                st.markdown("---")
                st.markdown("### 👁️ Report Preview")
                with st.expander("📄 View Report Content"):
                    st.markdown(f"**Case ID:** {case_id}")
                    st.markdown(f"**Language:** {case_language}")
                    st.markdown(f"**Date:** {datetime.now().strftime('%d-%m-%Y %H:%M')}")
                    st.markdown("---")
                    st.markdown("**Case Scenario:**")
                    st.markdown(case_scenario)
                    st.markdown("---")
                    st.markdown("**Analysis:**")
                    st.markdown(analysis[:500] + "..." if len(analysis) > 500 else analysis)
            else:
                st.error("❌ Case data not found!")
        else:
            st.info("📝 Analyze a case first to export reports!")
            st.markdown("""
            <div class="feature-card">
                <h3>📝 How to Export</h3>
                <ol>
                    <li>Go to "Analyze Case"</li>
                    <li>Enter case details and click "Analyze"</li>
                    <li>Return to this page</li>
                    <li>Choose Word (recommended) or PDF format</li>
                    <li>Download your professional report!</li>
                </ol>
                <br>
                <p><strong>💡 Pro Tip:</strong> Use Word format for Telugu/Hindi/Tamil - it displays all characters perfectly!</p>
            </div>
            """, unsafe_allow_html=True)
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #64748b; padding: 20px;">
        <p><strong>⚖️ LegalMitra</strong> - Your AI Legal Assistant | Powered by Google Gemini 2.5 Flash | 100% FREE</p>
        <p style="font-size: 14px;">🔒 Your data is secure | 🌐 Multi-language support | 💯 Professional Edition</p>
        <p style="font-size: 12px; margin-top: 10px;"><em>"Mitra" (मित्र) means "Friend" in Sanskrit</em></p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()